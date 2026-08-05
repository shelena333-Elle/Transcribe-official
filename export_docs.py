"""Экспорт структурированного текста в TXT / DOCX / PDF."""

from __future__ import annotations

import io
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from fpdf import FPDF


def _windows_font() -> str | None:
    candidates = [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _is_heading(line: str) -> bool:
    lower = line.lower()
    return (
        lower.startswith("блок")
        or lower.startswith("тематический")
        or lower == "диалог"
        or lower.startswith("основная")
    )


def _is_speaker(line: str) -> bool:
    return line.endswith(":") and len(line) < 60


def export_txt(text: str) -> bytes:
    return (text or "").encode("utf-8")


def export_docx(text: str, title: str = "Транскрипт") -> bytes:
    doc = Document()
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("=") and set(line) <= {"="}:
            continue

        if line and not line.endswith(":") and len(line) < 80 and _is_heading(line):
            p = doc.add_heading(line, level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x3A)
            continue

        if _is_speaker(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            continue

        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(text: str, title: str = "Транскрипт") -> bytes:
    font_path = _windows_font()
    pdf = FPDF(format="A4")
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    if font_path:
        pdf.add_font("Body", fname=font_path)
        pdf.add_font("Body", style="B", fname=font_path)
        family = "Body"
    else:
        family = "Helvetica"

    width = pdf.epw

    pdf.set_font(family, style="B", size=16)
    pdf.multi_cell(width, 9, title)
    pdf.ln(4)

    for raw_line in (text or "").splitlines():
        line = raw_line.rstrip()
        if not line:
            pdf.ln(3)
            continue
        if line.startswith("=") and set(line) <= {"="}:
            continue

        if _is_heading(line):
            pdf.ln(2)
            pdf.set_font(family, style="B", size=13)
            pdf.set_text_color(26, 58, 58)
            pdf.multi_cell(width, 8, line)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font(family, size=11)
            continue

        if _is_speaker(line):
            pdf.set_font(family, style="B", size=11)
            pdf.multi_cell(width, 7, line)
            pdf.set_font(family, size=11)
            continue

        pdf.set_font(family, size=11)
        if family == "Helvetica":
            line = line.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(width, 7, line)

    out = pdf.output()
    return bytes(out)


def export_bytes(text: str, fmt: str, title: str = "Транскрипт") -> tuple[bytes, str, str]:
    fmt = (fmt or "txt").lower().strip()
    if fmt in {"txt", "text"}:
        return export_txt(text), "text/plain; charset=utf-8", "transcription.txt"
    if fmt in {"docx", "doc", "word"}:
        return export_docx(text, title=title), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ), "transcription.docx"
    if fmt == "pdf":
        return export_pdf(text, title=title), "application/pdf", "transcription.pdf"
    raise ValueError(f"Неизвестный формат: {fmt}")
